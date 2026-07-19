from __future__ import annotations

import csv
import importlib.util
import json
import multiprocessing
import os
import queue
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable, Protocol

from src.observability import log_event


class ExtractBackend(Protocol):
    name: str

    def available(self) -> bool: ...

    def extract(self, file_path: str | Path) -> str | None:
        """Return extracted text; unavailable or failed backends return None."""


@dataclass(frozen=True)
class FileExtraction:
    file_name: str
    file_kind: str
    method_used: str
    rows: int
    needs_manual: bool
    note: str = ""
    confidence: float | None = None
    elapsed_seconds: float | None = None
    error_code: str | None = None


@dataclass(frozen=True)
class IngestionResult:
    extractions: list[FileExtraction]
    normalized: dict[str, list[dict[str, Any]]]
    needs_manual_count: int
    ok_count: int


class LLMVisionBackend:
    name = "llm_vision"

    def __init__(
        self,
        *,
        api_key_env: str = "CHAINGUARD_VISION_API_KEY",
        api_url_env: str = "CHAINGUARD_VISION_API_URL",
        timeout_seconds: float = 15.0,
    ) -> None:
        self.api_key_env = api_key_env
        self.api_url_env = api_url_env
        self.timeout_seconds = timeout_seconds

    def available(self) -> bool:
        try:
            return bool(os.environ.get(self.api_key_env, "").strip())
        except Exception:
            return False

    def extract(self, file_path: str | Path) -> str | None:
        try:
            api_key = os.environ.get(self.api_key_env, "").strip()
            api_url = os.environ.get(self.api_url_env, "").strip()
            if not api_key or not api_url:
                return None

            payload = json.dumps(
                {"file_name": Path(file_path).name},
                ensure_ascii=False,
            ).encode("utf-8")
            request = urllib.request.Request(
                api_url,
                data=payload,
                method="POST",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
            )
            with urllib.request.urlopen(request, timeout=self.timeout_seconds) as response:
                body = response.read().decode("utf-8")
            parsed = json.loads(body)
            text = parsed.get("text") or parsed.get("content") or parsed.get("result")
            return str(text).strip() if text else None
        except (
            OSError,
            TimeoutError,
            ValueError,
            urllib.error.URLError,
            urllib.error.HTTPError,
        ):
            return None
        except Exception:
            return None


class OcrEngineBackend:
    """Local Chinese/English OCR backed by RapidOCR and ONNX Runtime.

    The heavy engine is imported and initialized only inside a bounded child
    process.  API startup and non-image imports therefore do not load models,
    while a timed-out native inference can be terminated instead of lingering
    in the request worker.
    """

    name = "rapidocr_local"

    def __init__(
        self,
        *,
        timeout_seconds: float | None = None,
        min_confidence: float | None = None,
        max_pixels: int | None = None,
    ) -> None:
        self.timeout_seconds = timeout_seconds if timeout_seconds is not None else _env_float(
            "CHAINGUARD_OCR_TIMEOUT_SECONDS", 20.0, minimum=0.1
        )
        self.min_confidence = min_confidence if min_confidence is not None else _env_float(
            "CHAINGUARD_OCR_MIN_CONFIDENCE", 0.75, minimum=0.0, maximum=1.0
        )
        self.max_pixels = max_pixels if max_pixels is not None else _env_int(
            "CHAINGUARD_OCR_MAX_PIXELS", 25_000_000, minimum=1
        )
        self.last_metadata: dict[str, Any] = {}

    def available(self) -> bool:
        if os.environ.get("CHAINGUARD_OCR_ENABLED", "true").strip().lower() in {"0", "false", "no", "off"}:
            self.last_metadata = {
                "error_code": "OCR_DISABLED",
                "message": "本地 OCR 已通过配置禁用；需要人工处理。",
            }
            return False
        try:
            available = (
                importlib.util.find_spec("rapidocr") is not None
                and importlib.util.find_spec("onnxruntime") is not None
            )
        except Exception:
            return False
        if not available:
            self.last_metadata = {
                "error_code": "OCR_ENGINE_UNAVAILABLE",
                "message": "本地 RapidOCR/ONNX Runtime 不可用；需要人工处理。",
            }
        return available

    def extract(self, file_path: str | Path) -> str | None:
        started = time.perf_counter()
        path = Path(file_path)
        self._current_file_name = path.name
        self.last_metadata = {}
        try:
            from PIL import Image, UnidentifiedImageError

            with Image.open(path) as image:
                width, height = image.size
                image.verify()
            if width <= 0 or height <= 0 or width * height > self.max_pixels:
                return self._fail(
                    "OCR_IMAGE_LIMIT",
                    f"图片像素数超出 OCR 安全限制（上限 {self.max_pixels}）；需要人工处理。",
                    started,
                )
        except (OSError, ValueError, UnidentifiedImageError):
            return self._fail("OCR_IMAGE_DAMAGED", "图片损坏或格式无效；需要人工处理。", started)
        except Exception as error:
            return self._fail("OCR_IMAGE_DAMAGED", "图片无法安全解码；需要人工处理。", started, error)

        context = multiprocessing.get_context("spawn")
        result_queue = context.Queue(maxsize=1)
        process = context.Process(target=_rapidocr_worker, args=(str(path), result_queue), daemon=True)
        try:
            process.start()
            process.join(self.timeout_seconds)
            if process.is_alive():
                process.terminate()
                process.join(1.0)
                if process.is_alive() and hasattr(process, "kill"):
                    process.kill()
                    process.join(1.0)
                return self._fail(
                    "OCR_TIMEOUT",
                    f"本地 OCR 超过 {self.timeout_seconds:g} 秒超时；需要人工处理。",
                    started,
                )
            try:
                payload = result_queue.get(timeout=1.0)
            except queue.Empty:
                return self._fail("OCR_ENGINE_FAILED", "本地 OCR 未返回结果；需要人工处理。", started)
        except Exception as error:
            return self._fail("OCR_ENGINE_FAILED", "本地 OCR 启动失败；需要人工处理。", started, error)
        finally:
            result_queue.close()
            result_queue.join_thread()

        if payload.get("error"):
            return self._fail("OCR_ENGINE_FAILED", "本地 OCR 识别失败；需要人工处理。", started)
        texts = [str(value).strip() for value in payload.get("texts") or [] if str(value).strip()]
        scores = [float(value) for value in payload.get("scores") or []]
        if not texts or not scores:
            return self._fail("OCR_EMPTY", "图片为空白或未识别到文字；需要人工处理。", started)
        confidence = min(scores)
        elapsed = round(time.perf_counter() - started, 6)
        if confidence < self.min_confidence:
            self.last_metadata = {
                "confidence": round(confidence, 6),
                "elapsed_seconds": elapsed,
                "error_code": "OCR_LOW_CONFIDENCE",
                "message": (
                    f"OCR 最低文本置信度 {confidence:.3f} 低于阈值 "
                    f"{self.min_confidence:.3f}；需要人工处理。"
                ),
            }
            log_event(
                "ocr_manual_required", file_name=path.name, backend=self.name,
                error_code="OCR_LOW_CONFIDENCE", confidence=round(confidence, 6),
                elapsed_seconds=elapsed,
            )
            return None
        self.last_metadata = {
            "confidence": round(confidence, 6),
            "elapsed_seconds": elapsed,
            "line_count": len(texts),
        }
        log_event(
            "ocr_succeeded", file_name=path.name, backend=self.name,
            confidence=round(confidence, 6), line_count=len(texts), elapsed_seconds=elapsed,
        )
        return "\n".join(texts)

    def _fail(
        self,
        error_code: str,
        message: str,
        started: float,
        error: Exception | None = None,
    ) -> None:
        elapsed = round(time.perf_counter() - started, 6)
        self.last_metadata = {
            "elapsed_seconds": elapsed,
            "error_code": error_code,
            "message": message,
        }
        log_event(
            "ocr_manual_required", file_name=getattr(self, "_current_file_name", "-"), backend=self.name,
            error_code=error_code, exception=type(error).__name__ if error else None,
            elapsed_seconds=elapsed,
        )
        return None


def _rapidocr_worker(file_path: str, result_queue: Any) -> None:
    """Run native OCR in an isolated process; only pickle small safe results."""
    try:
        from rapidocr import RapidOCR  # type: ignore

        output = RapidOCR()(file_path, text_score=0.0)
        result_queue.put({
            "texts": list(getattr(output, "txts", ()) or ()),
            "scores": [float(value) for value in (getattr(output, "scores", ()) or ())],
        })
    except Exception as error:
        result_queue.put({"error": type(error).__name__})


def _env_float(name: str, default: float, *, minimum: float, maximum: float | None = None) -> float:
    try:
        value = float(os.environ.get(name, str(default)))
    except (TypeError, ValueError):
        return default
    value = max(value, minimum)
    return min(value, maximum) if maximum is not None else value


def _env_int(name: str, default: int, *, minimum: int) -> int:
    try:
        return max(int(os.environ.get(name, str(default))), minimum)
    except (TypeError, ValueError):
        return default


class TextLayerBackend:
    name = "text_layer"

    def available(self) -> bool:
        try:
            return (
                importlib.util.find_spec("pypdf") is not None
                or importlib.util.find_spec("pdfplumber") is not None
            )
        except Exception:
            return False

    def extract(self, file_path: str | Path) -> str | None:
        path = Path(file_path)
        if detect_kind(path) != "pdf":
            return None
        try:
            if importlib.util.find_spec("pypdf") is not None:
                from pypdf import PdfReader  # type: ignore

                reader = PdfReader(str(path))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                return text.strip() or None

            if importlib.util.find_spec("pdfplumber") is not None:
                import pdfplumber  # type: ignore

                with pdfplumber.open(str(path)) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                return text.strip() or None
        except Exception:
            return None
        return None


class WordTextBackend:
    """Extract paragraphs and tables from modern Word documents."""

    name = "word_text"

    def available(self) -> bool:
        try:
            return importlib.util.find_spec("docx") is not None
        except Exception:
            return False

    def extract(self, file_path: str | Path) -> str | None:
        path = Path(file_path)
        if path.suffix.lower() != ".docx":
            return None
        try:
            from docx import Document  # type: ignore

            document = Document(str(path))
            lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells]
                    if any(values):
                        lines.append("\t".join(values))
            return "\n".join(lines).strip() or None
        except Exception:
            return None


def detect_kind(file_path: str | Path) -> str:
    """Classify file kind by extension."""
    suffix = Path(file_path).suffix.lower()
    if suffix == ".csv":
        return "csv"
    if suffix in {".xlsx", ".xls"}:
        return "excel"
    if suffix == ".pdf":
        return "pdf"
    if suffix in {".doc", ".docx"}:
        return "word"
    if suffix in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}:
        return "image"
    return "unknown"


def default_backends() -> list[ExtractBackend]:
    """Build cascaded extraction backends with their own availability probes."""
    # Text-bearing PDF/DOCX remains cheap and deterministic. Images then use
    # local OCR first, with the optional remote backend only as a final fallback.
    return [TextLayerBackend(), WordTextBackend(), OcrEngineBackend(), LLMVisionBackend()]


def _extract_with_cascade_details(
    file_path: str | Path,
    backends: list[ExtractBackend] | None = None,
) -> tuple[str | None, str, dict[str, Any]]:
    last_metadata: dict[str, Any] = {}
    for backend in backends or default_backends():
        try:
            if not backend.available():
                metadata = getattr(backend, "last_metadata", None)
                if metadata:
                    last_metadata = dict(metadata)
                continue
        except Exception:
            continue
        try:
            text = backend.extract(file_path)
        except Exception:
            text = None
        metadata = getattr(backend, "last_metadata", None)
        if metadata:
            last_metadata = dict(metadata)
        if text:
            return text, backend.name, last_metadata
    return None, "manual_required", last_metadata


def extract_with_cascade(
    file_path: str | Path,
    backends: list[ExtractBackend] | None = None,
) -> tuple[str | None, str]:
    """Try available backends in order and return extracted text plus method."""
    text, method, _ = _extract_with_cascade_details(file_path, backends)
    return text, method


def ingest_files(
    file_paths: list[str | Path],
    *,
    backends: list[ExtractBackend] | None = None,
    csv_reader: Callable[[str | Path], list[dict[str, Any]]] | None = None,
) -> IngestionResult:
    """Detect, extract, normalize, and summarize a batch of source files."""
    normalized: dict[str, list[dict[str, Any]]] = {}
    extractions: list[FileExtraction] = []
    reader = csv_reader or _read_csv_rows

    for file_path in file_paths:
        path = Path(file_path)
        kind = detect_kind(path)
        stem = path.stem.lower()

        if kind == "csv":
            rows, note = _safe_read_csv(reader, path)
            table_name = stem
            if rows:
                normalized[table_name] = rows
                extractions.append(
                    FileExtraction(path.name, kind, "direct", len(rows), False)
                )
            else:
                extractions.append(
                    FileExtraction(
                        path.name,
                        kind,
                        "direct",
                        0,
                        True,
                        note or "No rows extracted; manual review required",
                    )
                )
            continue

        if kind in {"excel", "pdf", "image", "word"}:
            text, method_used, metadata = _extract_with_cascade_details(path, backends)
            rows = _text_to_rows(text or "")
            if rows:
                normalized[f"{stem}_intake"] = rows
                extractions.append(
                    FileExtraction(
                        path.name, kind, method_used, len(rows), False,
                        confidence=metadata.get("confidence"),
                        elapsed_seconds=metadata.get("elapsed_seconds"),
                    )
                )
            else:
                extractions.append(
                    FileExtraction(
                        path.name,
                        kind,
                        "manual_required",
                        0,
                        True,
                        str(metadata.get("message") or "未识别到可导入内容；需要人工处理。"),
                        confidence=metadata.get("confidence"),
                        elapsed_seconds=metadata.get("elapsed_seconds"),
                        error_code=metadata.get("error_code") or "EXTRACTION_EMPTY",
                    )
                )
            continue

        extractions.append(
            FileExtraction(
                path.name,
                kind,
                "manual_required",
                0,
                True,
                "Unsupported file type; manual entry required",
            )
        )

    needs_manual_count = sum(1 for extraction in extractions if extraction.needs_manual)
    return IngestionResult(
        extractions=extractions,
        normalized=normalized,
        needs_manual_count=needs_manual_count,
        ok_count=len(extractions) - needs_manual_count,
    )


def _read_csv_rows(file_path: str | Path) -> list[dict[str, Any]]:
    with Path(file_path).open(newline="", encoding="utf-8-sig") as handle:
        return [dict(row) for row in csv.DictReader(handle)]


def _safe_read_csv(
    reader: Callable[[str | Path], list[dict[str, Any]]],
    path: Path,
) -> tuple[list[dict[str, Any]], str]:
    try:
        return reader(path), ""
    except Exception as exc:
        return [], f"CSV parse failed: {exc}"


def _text_to_rows(text: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        values = _split_text_line(line)
        if not values:
            continue
        rows.append({f"col_{index + 1}": value for index, value in enumerate(values)})
    return rows


def _split_text_line(line: str) -> list[str]:
    if "\t" in line:
        return [part.strip() for part in line.split("\t")]
    try:
        parsed = next(csv.reader([line]))
        return [part.strip() for part in parsed]
    except Exception:
        return [line]
